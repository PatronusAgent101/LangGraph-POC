from docx import Document
from datetime import datetime
import json

def extract_docx_content_by_pages_2(file_path):
    doc = Document(file_path)

    metadata = {
        "author": doc.core_properties.author or "Unknown",
        "created": doc.core_properties.created.isoformat() if doc.core_properties.created else datetime.now().isoformat()
    }

    pages = []
    page_count = 1
    current_content = {
        "heading1": None,
        "heading2": None,
        "paragraph": [],
        "table": [],
        "footers": []
    }

    def start_new_page():
        nonlocal page_count, current_content
        pages.append({
            "page": str(page_count),
            "content": current_content
        })
        page_count += 1
        current_content = {
            "heading1": None,
            "heading2": None,
            "paragraph": [],
            "table": [],
            "footers": []
        }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Look for headings
        style = para.style.name.lower()
        if "heading 1" in style:
            current_content["heading1"] = text
        elif "heading 2" in style:
            current_content["heading2"] = text
        else:
            # Check for manual page break in XML
            run_xmls = [run._element.xml for run in para.runs]
            if any("<w:br w:type=\"page\"" in xml for xml in run_xmls):
                start_new_page()
            else:
                current_content["paragraph"].append(text)

        # Heuristic: flush page every ~10 paragraphs
        if len(current_content["paragraph"]) >= 10:
            start_new_page()

    # After all paragraphs, push remaining content
    if any(current_content.values()):
        pages.append({
            "page": str(page_count),
            "content": current_content
        })

    # Tables go into the last page or create new if empty
    if doc.tables:
        if not pages:
            pages.append({
                "page": str(page_count),
                "content": {
                    "heading1": None,
                    "heading2": None,
                    "paragraph": [],
                    "table": [],
                    "footers": []
                }
            })
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            pages[-1]["content"]["table"].append(rows)

    # Footer from first section
    if doc.sections:
        footer_texts = doc.sections[0].footer.paragraphs
        if footer_texts:
            pages[0]["content"]["footers"].append(" ".join(p.text.strip() for p in footer_texts if p.text.strip()))

    result = {
        "metadata": metadata,
        "structure": {
            "pages": pages
        }
    }

    return result
   # Example usage:
file_path = "Professional Summary Draft.docx"  # Update this path
json_output = extract_docx_content_by_pages_2(file_path)

# Print or save the output
print(json.dumps(json_output, indent=4))
